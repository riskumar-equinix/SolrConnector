#!/usr/bin/env python3
"""Convert Markdown to Word docx using python-docx."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def parse_table(lines):
    """Parse markdown table into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if all(re.match(r"^[-:]+$", c) for c in cells):
                continue
            rows.append(cells)
    return rows


def parse_markdown(content):
    """Parse markdown content into structured blocks."""
    blocks = []
    lines = content.split("\n")
    i = 0
    in_table = False
    table_lines = []
    in_code = False
    code_lines = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()

        if line_stripped.startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(code_lines), code_lang))
                code_lines = []
                in_code = False
            else:
                code_lang = line_stripped[3:].strip()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line_stripped.startswith("|") and "|" in line_stripped[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                rows = parse_table(table_lines)
                if rows:
                    blocks.append(("table", rows))
                in_table = False
                table_lines = []

        if line_stripped == "---":
            blocks.append(("hr",))
            i += 1
            continue

        if line_stripped.startswith("# "):
            blocks.append(("h1", line_stripped[2:].strip()))
            i += 1
            continue
        if line_stripped.startswith("## "):
            blocks.append(("h2", line_stripped[3:].strip()))
            i += 1
            continue
        if line_stripped.startswith("### "):
            blocks.append(("h3", line_stripped[4:].strip()))
            i += 1
            continue

        if line_stripped.startswith("- "):
            blocks.append(("list", line_stripped[2:].strip()))
            i += 1
            continue

        if line_stripped.startswith("*") and line_stripped.endswith("*") and not line_stripped.startswith("**"):
            blocks.append(("list", line_stripped[1:-1].strip()))
            i += 1
            continue

        if line_stripped:
            blocks.append(("p", line_stripped))
        else:
            blocks.append(("blank",))
        i += 1

    if in_table and table_lines:
        rows = parse_table(table_lines)
        if rows:
            blocks.append(("table", rows))
    if in_code and code_lines:
        blocks.append(("code", "\n".join(code_lines), code_lang))

    return blocks


def add_formatted_paragraph(doc, text):
    """Add paragraph with bold/italic handling."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2] + " ")
            run.bold = True
        else:
            p.add_run(part)


def md_to_docx(md_path, docx_path):
    """Convert markdown file to docx."""
    content = Path(md_path).read_text(encoding="utf-8")
    blocks = parse_markdown(content)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    for block in blocks:
        if block[0] == "h1":
            p = doc.add_heading(block[1], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif block[0] == "h2":
            doc.add_heading(block[1], level=1)
        elif block[0] == "h3":
            doc.add_heading(block[1], level=2)
        elif block[0] == "p":
            add_formatted_paragraph(doc, block[1])
        elif block[0] == "list":
            add_formatted_paragraph(doc, "• " + block[1])
        elif block[0] == "table":
            rows = block[1]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
        elif block[0] == "code":
            p = doc.add_paragraph()
            p.style = "Normal"
            run = p.add_run(block[1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif block[0] == "hr":
            doc.add_paragraph("_" * 60)
        elif block[0] == "blank":
            doc.add_paragraph()

    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        md_file = Path(__file__).parent.parent / "docs" / "SDD - Global Search Data Ingestion and Search.docx.md"
        docx_file = md_file.parent / md_file.name.replace(".docx.md", ".docx").replace(".md", ".docx")
    else:
        md_file = Path(sys.argv[1])
        docx_file = Path(sys.argv[2]) if len(sys.argv) > 2 else md_file.with_suffix(".docx")
    md_to_docx(md_file, docx_file)
